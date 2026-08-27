#!/usr/bin/env python3
"""
Convert Implementation Plan Markdown to a beautifully styled .docx document
for Google Docs and Microsoft Word.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Styles
    normal_style = doc.styles['Normal']
    font = normal_style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(32, 33, 36)

    # Document Title
    p_title = doc.add_paragraph()
    run_title = p_title.add_run("GCP Production Data Ingestion Architecture Plan (v3.0)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(26, 115, 232)  # Google Blue
    p_title.space_after = Pt(4)

    p_sub = doc.add_paragraph()
    run_sub = p_sub.add_run("GCS to BigQuery Batch Pipeline with End-to-End Testing & Deployment Framework")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(95, 99, 104)
    p_sub.space_after = Pt(24)

    # Read Markdown source
    md_file = "/usr/local/google/home/pradeepsarathy/.gemini/antigravity/brain/c43cd83c-9ed4-45a0-8f41-4a377feee68a/implementation_plan.md"
    with open(md_file, "r", encoding="utf-8") as f:
        lines = f.readlines()

    in_table = False
    table_rows = []

    for line in lines:
        line_str = line.strip()
        if not line_str or line_str.startswith("# GCP Production") or line_str.startswith("**GCS to BigQuery") or line_str == "---":
            continue

        # Section Headings
        if line_str.startswith("## "):
            h_text = line_str[3:].strip()
            p = doc.add_paragraph()
            r = p.add_run(h_text)
            r.font.size = Pt(16)
            r.font.bold = True
            r.font.color.rgb = RGBColor(30, 142, 62)  # Google Green
            p.space_before = Pt(18)
            p.space_after = Pt(8)
            continue
        elif line_str.startswith("### "):
            h_text = line_str[4:].strip()
            p = doc.add_paragraph()
            r = p.add_run(h_text)
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(217, 48, 37)  # Google Red
            p.space_before = Pt(14)
            p.space_after = Pt(6)
            continue

        # Tables
        if "|" in line_str and ("---" in line_str or line_str.startswith("|")):
            if "---" in line_str:
                continue
            cells = [c.strip() for c in line_str.split("|")[1:-1]]
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            # Build Table
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, text in enumerate(row_data):
                        cell = table.cell(row_idx, col_idx)
                        cell.text = text
                        set_cell_margins(cell)
                        p_cell = cell.paragraphs[0]
                        p_cell.paragraph_format.space_before = Pt(2)
                        p_cell.paragraph_format.space_after = Pt(2)

                        if row_idx == 0:
                            set_cell_background(cell, "F1F3F4")
                            for run in p_cell.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(32, 33, 36)
                        else:
                            if row_idx % 2 == 0:
                                set_cell_background(cell, "F8F9FA")

            table_rows = []
            in_table = False

        # Code Blocks / Mermaid / Lists
        if line_str.startswith("```"):
            continue

        if line_str.startswith("* ") or line_str.startswith("- "):
            item_text = line_str[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Bold parsing
            parts = item_text.split("**")
            for i, part in enumerate(parts):
                r = p.add_run(part)
                if i % 2 == 1:
                    r.font.bold = True
            continue

        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        parts = line_str.split("**")
        for i, part in enumerate(parts):
            r = p.add_run(part)
            if i % 2 == 1:
                r.font.bold = True

    output_path = "/usr/local/google/home/pradeepsarathy/AntiGravity_Projects/Project_3/coned_demo/GCP_Data_Ingestion_Architecture_Plan_v3.docx"
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_document()
